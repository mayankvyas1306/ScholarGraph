import os
import logging
import docx
from typing import List, Dict, Any
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from backend.data.models import PaperMeta, GapClaim, Summary
from backend.clients.claude_client import ClaudeClient

logger = logging.getLogger("researchmind.report")

def generate_introduction(query: str, summaries: List[Summary], gap_claims: List[GapClaim]) -> str:
    """
    Uses Gemini to write a rich, multi-paragraph introduction for the report.
    """
    logger.info("Generating report introduction...")
    gemini = ClaudeClient()
    num_papers = len(summaries)
    num_gaps   = len(gap_claims)

    prompt = f"""
Write a formal, publication-quality Introduction section for an automated academic literature review report.

Report details:
- Research topic: "{query}"
- Number of papers analysed: {num_papers}
- Number of research gaps identified: {num_gaps}
- Data sources: arXiv and Semantic Scholar

The introduction MUST cover all of the following points in flowing prose (no bullet lists):
1. Briefly motivate WHY the topic '{query}' is important and timely in the current research landscape.
2. State the objectives of this review — what questions it seeks to answer.
3. Describe the methodology at a high level: automated corpus collection ({num_papers} papers), structured field extraction (method, dataset, key metric, limitation), citation-graph analysis, and thematic synthesis.
4. Mention that {num_gaps} candidate research gaps were detected and will be discussed.
5. Outline the structure of the report (Comparison Matrix → Thematic Synthesis → Research Gaps).

Length: 5–7 well-developed sentences organized into 2 short paragraphs.
Style: formal third-person academic prose, no markdown headers, no bullet points.

Return ONLY the plain prose text.
"""
    try:
        text = gemini.complete(
            prompt=prompt,
            system="You are an expert academic research writer producing publication-grade report introductions.",
            max_tokens=600,
            temperature=0.2,
        ).strip()
        return text
    except Exception as exc:
        logger.error(f"Failed to generate introduction: {exc}")
        # Compact deterministic fallback
        return (
            f"This report presents an automated systematic literature survey and candidate research gap discovery "
            f"for the topic '{query}'. A corpus of {num_papers} relevant papers was retrieved from arXiv and "
            f"Semantic Scholar, after which structured methodology records were extracted for each work covering "
            f"proposed methods, evaluation datasets, key performance metrics, and acknowledged limitations. "
            f"A citation-graph analysis was subsequently performed to surface {num_gaps} candidate research gaps "
            f"that current literature has not yet fully addressed. "
            f"The remainder of this report is organised as follows: Section 2 presents a side-by-side Comparison "
            f"Matrix of all analysed papers; Section 3 provides a thematic synthesis grouping works by methodology "
            f"and evaluation approach; and Section 4 details the identified research gaps with suggested future directions."
        )


def generate_thematic_synthesis(query: str, summaries: List[Summary], gap_claims: List[GapClaim]) -> str:
    """
    Leverages Gemini to synthesize a deep, publication-grade academic survey of the literature.
    """
    logger.info("Generating thematic academic synthesis of the literature...")
    gemini = ClaudeClient()

    papers_input = ""
    for idx, s in enumerate(summaries):
        papers_input += f"Paper {idx+1}: {s.title}\nSummary & Methodology: {s.summary_text}\n\n"

    gaps_input = ""
    for idx, gap in enumerate(gap_claims):
        gaps_input += (
            f"Gap {idx+1}: {gap.topic_label}\n"
            f"Description: {gap.description}\n"
            f"Suggested Directions: {', '.join(gap.suggested_directions)}\n\n"
        )

    prompt = f"""
Write a deep, professional, publication-grade thematic literature review and synthesis on: '{query}'.

You have {len(summaries)} papers to synthesize. Use ALL of them — do not skip any.

Paper summaries and extracted methodology records:
{papers_input}

Identified literature gaps to weave into the discussion:
{gaps_input}

STRICT WRITING REQUIREMENTS:
1. Academic third-person style throughout — no first-person, no casual language.
2. Organize into EXACTLY FOUR numbered subsections:
   ### 3.1 Methodological Landscape
      - Group papers by approach/architecture family. Compare and contrast techniques.
      - Highlight how methods evolved over time (use years where possible).
      - Include specific model names, algorithmic details from the summaries.
   ### 3.2 Datasets, Benchmarks & Evaluation Protocols
      - Name the actual datasets/benchmarks used across the papers.
      - Compare evaluation metrics (accuracy, BLEU, F1, mAP, perplexity, etc.).
      - Note any inconsistencies or lack of standardization in evaluation.
   ### 3.3 Common Limitations & Open Challenges
      - Synthesize the recurring limitations found across papers.
      - Explain how these limitations create the research gaps identified.
      - Be specific — avoid vague statements like "more research is needed".
   ### 3.4 Synthesis & Critical Assessment
      - Provide your overall critical assessment of the field's maturity.
      - Identify which directions appear most promising and why.
      - Connect clearly to the {len(gap_claims)} identified research gaps.
3. Cite papers inline using author surname or short title (e.g., [Smith et al.], [PyramidTNT]).
   Do NOT simply list papers sequentially — group, compare, and contrast them.
4. Include quantitative values (accuracy scores, parameter counts, dataset sizes, years) wherever the summaries provide them.
5. Each subsection must be at least 3 substantial paragraphs.
6. Total length: 1000–1400 words. A reader unfamiliar with the topic should come away with a clear mental model of the field.

Return ONLY the synthesized text with the ### subsection headers. No markdown code fences, no preamble.
"""
    try:
        response = gemini.complete(
            prompt=prompt,
            system=(
                "You are a senior academic research writer and expert literature review synthesiser. "
                "You write with precision, depth, and clarity — producing content that would be "
                "accepted in a top-tier journal survey paper."
            ),
            max_tokens=4000,
            temperature=0.2,
        )
        return response.strip()
    except Exception as e:
        logger.error(f"Failed to generate thematic synthesis: {e}")
        fallback = "### 3.1 Literature Survey Overview\n"
        for s in summaries:
            fallback += f"**{s.title}**: {s.summary_text}\n\n"
        return fallback


def generate_gap_narratives(query: str, gap_claims: List[GapClaim], summaries: List[Summary]) -> List[str]:
    """
    Uses Gemini to write a rich explanatory narrative paragraph for each research gap.
    Returns a list of narrative strings, one per gap_claim.
    """
    if not gap_claims:
        return []

    logger.info(f"Generating narrative paragraphs for {len(gap_claims)} research gaps...")
    gemini = ClaudeClient()

    paper_titles = [s.title for s in summaries]

    narratives = []
    for gap in gap_claims:
        directions_text = "\n".join(f"- {d}" for d in gap.suggested_directions)
        prompt = f"""
You are writing Section 4 of an academic literature review report on '{query}'.

Write a rich, explanatory narrative paragraph (120–180 words) for the following identified research gap.

Gap topic: {gap.topic_label}
Gap description: {gap.description}
Citation density: {gap.citation_density:.2f} citations/paper (lower = more underexplored)
Suggested future directions:
{directions_text}

The broader paper corpus covered these works:
{', '.join(paper_titles[:20])}

Your paragraph MUST:
1. Explain WHY this gap exists — what about current methods or evaluation leaves it unaddressed.
2. Describe the SIGNIFICANCE of closing this gap for the field of '{query}'.
3. Connect at least one of the suggested directions to a concrete research methodology (e.g., a specific model family, evaluation protocol, or dataset type that could be used).
4. Be written in formal third-person academic prose — no bullet points, no headings.
5. Flow naturally after a heading like "Gap N: {gap.topic_label}".

Return ONLY the plain narrative paragraph text.
"""
        try:
            text = gemini.complete(
                prompt=prompt,
                system="You are an expert academic research writer specializing in research gap analysis.",
                max_tokens=400,
                temperature=0.2,
            ).strip()
            narratives.append(text)
        except Exception as exc:
            logger.error(f"Failed to generate gap narrative for '{gap.topic_label}': {exc}")
            narratives.append(gap.description)

    return narratives

def add_markdown_paragraphs_docx(doc, text: str):
    """
    Parses a simple markdown text block and adds it to python-docx Document.
    """
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)

def add_markdown_paragraphs_pdf(story, text: str, h1_style, h2_style, body_style):
    """
    Parses a simple markdown text block and adds it to ReportLab story.
    """
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("### "):
            story.append(Paragraph(line[4:], h2_style))
            story.append(Spacer(1, 4))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:], h1_style))
            story.append(Spacer(1, 6))
        elif line.startswith("- "):
            story.append(Paragraph(f"• {line[2:]}", body_style))
            story.append(Spacer(1, 4))
        else:
            story.append(Paragraph(line, body_style))
            story.append(Spacer(1, 6))

def compile_markdown_draft(
    query: str,
    summaries: List[Summary],
    comparison_table: List[Dict[str, Any]],
    gap_claims: List[GapClaim],
    synthesis_text: str,
    introduction_text: str = "",
    gap_narratives: List[str] = None,
) -> str:
    """
    Compiles a plain text/markdown draft of the final report.
    """
    if gap_narratives is None:
        gap_narratives = []

    intro = (
        f"# ResearchMind Literature Review Report\n\n"
        f"**Research Topic:** {query}\n\n"
        f"## 1. Introduction\n\n"
        f"{introduction_text}\n"
    )

    table_section = "\n## 2. Comparison Matrix\n\n"
    table_section += "| Title | Year | Method | Dataset | Key Metric | Limitation |\n"
    table_section += "|---|---|---|---|---|---|\n"
    for row in comparison_table:
        title_trunc = row['title'][:40] + "..." if len(row['title']) > 40 else row['title']
        table_section += (
            f"| {title_trunc} | {row['year']} | {row['method']} "
            f"| {row['dataset']} | {row.get('key_metric', 'N/A')} | {row['limitation']} |\n"
        )

    summary_section = f"\n## 3. Thematic Literature Survey & Synthesis\n\n{synthesis_text}\n"

    gap_section = "\n## 4. Identified Research Gaps\n\n"
    if not gap_claims:
        gap_section += "No high-confidence research gaps were detected in this literature set.\n"
    else:
        for idx, gap in enumerate(gap_claims):
            gap_section += f"### Gap {idx+1}: {gap.topic_label}\n\n"
            narrative = gap_narratives[idx] if idx < len(gap_narratives) else gap.description
            gap_section += f"{narrative}\n\n"
            gap_section += f"**Citation Density:** {gap.citation_density:.2f} citations/paper\n\n"
            gap_section += "**Suggested Directions for Future Research:**\n"
            for dir_stmt in gap.suggested_directions:
                gap_section += f"- {dir_stmt}\n"
            gap_section += "\n"

    return intro + table_section + summary_section + gap_section

def generate_docx(
    output_path: str,
    query: str,
    summaries: List[Summary],
    comparison_table: List[Dict[str, Any]],
    gap_claims: List[GapClaim],
    synthesis_text: str,
    introduction_text: str = "",
    gap_narratives: List[str] = None,
):
    """
    Generates a structured DOCX report.
    """
    if gap_narratives is None:
        gap_narratives = []

    doc = docx.Document()

    # Title
    doc.add_heading("ResearchMind Literature Review & Gap Analysis", 0)
    doc.add_paragraph(f"Focus Topic: {query}").bold = True

    # Section 1: Introduction
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(introduction_text or (
        f"This report presents an automated literature review on '{query}'. "
        f"{len(summaries)} papers were analysed from arXiv and Semantic Scholar."
    ))

    # Section 2: Comparison Table
    doc.add_heading("2. Comparison Matrix", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Title'
    hdr_cells[1].text = 'Year'
    hdr_cells[2].text = 'Method'
    hdr_cells[3].text = 'Dataset'
    hdr_cells[4].text = 'Limitation'

    for row in comparison_table:
        row_cells = table.add_row().cells
        row_cells[0].text = row['title'][:30] + "..." if len(row['title']) > 30 else row['title']
        row_cells[1].text = str(row['year'])
        row_cells[2].text = row['method']
        row_cells[3].text = row['dataset']
        row_cells[4].text = row['limitation']

    # Section 3: Thematic Synthesis
    doc.add_heading("3. Thematic Literature Survey & Synthesis", level=1)
    add_markdown_paragraphs_docx(doc, synthesis_text)

    # Section 4: Research Gaps
    doc.add_heading("4. Identified Research Gaps", level=1)
    if not gap_claims:
        doc.add_paragraph("No significant research gaps were identified in the corpus.")
    else:
        for idx, gap in enumerate(gap_claims):
            doc.add_heading(f"Gap {idx+1}: {gap.topic_label}", level=2)
            narrative = gap_narratives[idx] if idx < len(gap_narratives) else gap.description
            doc.add_paragraph(narrative)
            doc.add_paragraph(f"Citation Density: {gap.citation_density:.2f} citations/paper")
            p = doc.add_paragraph("Suggested Future Directions:")
            p.bold = True
            for dir_stmt in gap.suggested_directions:
                doc.add_paragraph(dir_stmt, style='List Bullet')

    doc.save(output_path)
    logger.info(f"DOCX report saved to {output_path}")

def generate_pdf(
    output_path: str,
    query: str,
    summaries: List[Summary],
    comparison_table: List[Dict[str, Any]],
    gap_claims: List[GapClaim],
    synthesis_text: str,
    introduction_text: str = "",
    gap_narratives: List[str] = None,
):
    """
    Generates a premium PDF report using ReportLab.
    """
    if gap_narratives is None:
        gap_narratives = []

    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()

    # ── Custom Paragraph Styles ──────────────────────────────────────────────
    title_style = ParagraphStyle(
        'DocTitle', parent=styles['Title'],
        fontName='Helvetica-Bold', fontSize=24, leading=28,
        textColor=colors.HexColor('#1E293B'), spaceAfter=15,
    )
    h1_style = ParagraphStyle(
        'Heading1Style', parent=styles['Heading1'],
        fontName='Helvetica-Bold', fontSize=16, leading=20,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=14, spaceAfter=8, keepWithNext=True,
    )
    h2_style = ParagraphStyle(
        'Heading2Style', parent=styles['Heading2'],
        fontName='Helvetica-Bold', fontSize=12, leading=16,
        textColor=colors.HexColor('#1E293B'),
        spaceBefore=10, spaceAfter=4, keepWithNext=True,
    )
    body_style = ParagraphStyle(
        'BodyStyle', parent=styles['Normal'],
        fontName='Helvetica', fontSize=10, leading=15,
        textColor=colors.HexColor('#334155'), spaceAfter=8,
    )
    table_hdr_style = ParagraphStyle(
        'TableHdrStyle', parent=styles['Normal'],
        fontName='Helvetica-Bold', fontSize=8, leading=10,
        textColor=colors.white,
    )
    table_cell_style = ParagraphStyle(
        'TableCellStyle', parent=styles['Normal'],
        fontName='Helvetica', fontSize=8, leading=10,
        textColor=colors.HexColor('#334155'),
    )

    story = []

    # ── Cover ────────────────────────────────────────────────────────────────
    story.append(Paragraph("Literature Review &amp; Gap Analysis Report", title_style))
    story.append(Paragraph(f"<b>Research Focus:</b> {query}", body_style))
    story.append(Spacer(1, 10))

    # ── Section 1: Introduction ──────────────────────────────────────────────
    story.append(Paragraph("1. Introduction", h1_style))
    intro_text = introduction_text or (
        f"This report presents an automated systematic literature survey for the query '{query}'. "
        f"A corpus of {len(summaries)} relevant papers was retrieved, parsed, and synthesised "
        f"to analyse the current state of methodology in this domain."
    )
    # Split intro into paragraphs on double newlines for better readability
    for para in intro_text.split("\n\n"):
        para = para.strip()
        if para:
            story.append(Paragraph(para, body_style))
    story.append(Spacer(1, 10))

    # ── Section 2: Comparison Table ──────────────────────────────────────────
    story.append(Paragraph("2. Comparison Matrix", h1_style))
    data = [[
        Paragraph("Title", table_hdr_style),
        Paragraph("Year", table_hdr_style),
        Paragraph("Method", table_hdr_style),
        Paragraph("Dataset", table_hdr_style),
        Paragraph("Limitation", table_hdr_style),
    ]]
    for row in comparison_table:
        title_short = row['title'][:28] + "..." if len(row['title']) > 28 else row['title']
        data.append([
            Paragraph(title_short, table_cell_style),
            Paragraph(str(row['year']), table_cell_style),
            Paragraph(row['method'], table_cell_style),
            Paragraph(row['dataset'], table_cell_style),
            Paragraph(row['limitation'], table_cell_style),
        ])
    t = Table(data, colWidths=[145, 35, 110, 110, 115])
    t.setStyle(TableStyle([
        ('BACKGROUND',    (0, 0), (-1, 0),  colors.HexColor('#1E293B')),
        ('ALIGN',         (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING',    (0, 0), (-1, -1), 6),
        ('GRID',          (0, 0), (-1, -1), 0.5, colors.HexColor('#E2E8F0')),
        ('BACKGROUND',    (0, 1), (-1, -1), colors.HexColor('#F8FAFC')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#F8FAFC'), colors.white]),
    ]))
    story.append(t)
    story.append(Spacer(1, 15))
    story.append(PageBreak())

    # ── Section 3: Thematic Synthesis ────────────────────────────────────────
    story.append(Paragraph("3. Thematic Literature Survey &amp; Synthesis", h1_style))
    add_markdown_paragraphs_pdf(story, synthesis_text, h1_style, h2_style, body_style)
    story.append(Spacer(1, 10))
    story.append(PageBreak())

    # ── Section 4: Research Gaps ─────────────────────────────────────────────
    story.append(Paragraph("4. Identified Research Gaps", h1_style))
    if not gap_claims:
        story.append(Paragraph("No significant research gaps were identified in the corpus.", body_style))
    else:
        for idx, gap in enumerate(gap_claims):
            story.append(Paragraph(f"Gap {idx + 1}: {gap.topic_label}", h2_style))
            # Rich narrative paragraph
            narrative = gap_narratives[idx] if idx < len(gap_narratives) else gap.description
            for para in narrative.split("\n\n"):
                para = para.strip()
                if para:
                    story.append(Paragraph(para, body_style))
            story.append(Paragraph(
                f"<b>Citation Density:</b> {gap.citation_density:.2f} citations/paper",
                body_style,
            ))
            story.append(Paragraph("<b>Suggested Future Directions:</b>", body_style))
            for dir_stmt in gap.suggested_directions:
                story.append(Paragraph(f"&bull; {dir_stmt}", body_style))
            story.append(Spacer(1, 12))

    doc.build(story)
    logger.info(f"PDF report saved to {output_path}")

def run_report(state: dict) -> dict:
    """
    Compiles summaries and gap claims into PDF and DOCX documents.
    Generates Gemini-written Introduction, Thematic Synthesis, and
    per-gap narrative paragraphs before assembling the final documents.
    """
    query = state.get("query", "")
    summaries: List[Summary] = state.get("summaries", [])
    comparison_table: List[Dict[str, Any]] = state.get("comparison_table", [])
    gap_claims: List[GapClaim] = state.get("gap_claims", [])

    if "agent_status" not in state:
        state["agent_status"] = {}

    state["agent_status"]["report"] = "running"
    logger.info("Report Agent: Commencing report compilation.")

    # 1. Generate all three LLM-written sections
    logger.info("Step 1/3 — Generating Introduction...")
    introduction_text = generate_introduction(query, summaries, gap_claims)

    logger.info("Step 2/3 — Generating Thematic Synthesis...")
    synthesis_text = generate_thematic_synthesis(query, summaries, gap_claims)

    logger.info("Step 3/3 — Generating Research Gap narratives...")
    gap_narratives = generate_gap_narratives(query, gap_claims, summaries)

    # 2. Compile markdown draft (stored in state for frontend preview)
    state["report_draft"] = {
        "text": compile_markdown_draft(
            query, summaries, comparison_table, gap_claims,
            synthesis_text,
            introduction_text=introduction_text,
            gap_narratives=gap_narratives,
        ),
        "synthesis_text": synthesis_text,
        "introduction_text": introduction_text,
    }

    # 3. Write output files
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    exports_dir = os.path.join(base_dir, "db", "exports")
    os.makedirs(exports_dir, exist_ok=True)

    pdf_path  = os.path.join(exports_dir, "report.pdf")
    docx_path = os.path.join(exports_dir, "report.docx")

    try:
        generate_docx(
            docx_path, query, summaries, comparison_table, gap_claims,
            synthesis_text,
            introduction_text=introduction_text,
            gap_narratives=gap_narratives,
        )
        state["report_draft"]["docx_path"] = docx_path
    except Exception as e:
        logger.error(f"Failed to generate DOCX report: {e}")

    try:
        generate_pdf(
            pdf_path, query, summaries, comparison_table, gap_claims,
            synthesis_text,
            introduction_text=introduction_text,
            gap_narratives=gap_narratives,
        )
        state["report_draft"]["pdf_path"] = pdf_path
    except Exception as e:
        logger.error(f"Failed to generate PDF report: {e}")

    state["agent_status"]["report"] = "done"
    logger.info("Report Agent: All sections generated successfully.")
    return state
